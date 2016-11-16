# -*- coding:utf-8 -*-
import urllib.request
import os
import shutil
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import codecs

def get_html_soup(url,code):#获取解编码后的HTML
    html = None
    try:
        headers = {'User-Agent':'Mozilla/5.0 (Windows; U; Windows NT 6.1; en-US; rv:1.9.1.6) Gecko/20091201 Firefox/3.5.6'}
        req = urllib.request.Request(url=url, headers=headers)
        html = urllib.request.urlopen(req).read().decode(encoding = code, errors='ignore')
    except Exception as e:
        # print(e, "please check your network situation")
        return None
    soup = BeautifulSoup(str(html), "lxml")
    
    return soup

def getClasslink(url,pattern):
    soup=get_html_soup(url,'gbk')
    class_link={}
    scroll_list = BeautifulSoup(str(soup.find("div", attrs = pattern)), "lxml")
    for link in scroll_list.find_all("a"):
        if len(link.get_text().strip()) > 0 and link.get("href").find("http") != -1:
            if link.get('href').find("bbs")==-1 and link.get('href').find("960")==-1 and link.get('href').find("photo")==-1 and \
            link.get('href').find("house") ==-1 and link.get('href').find("v.yangtse")==-1:
                class_link[link.get_text()] = link.get('href')
            
    return class_link
    
def get_title_link(url, pattern, totalpage):#获取新闻的标题和正文链接
    news_link = {}
    for i in range(1, totalpage):
        numurl=page_url(url, i)
        soup = get_html_soup(numurl,'gbk')
        if soup==None:
            continue
        scroll_list = BeautifulSoup(str(soup.find("ul", attrs = pattern)), "lxml")
        for link in scroll_list.find_all("a"):
            if len(link.get_text().strip()) > 0 and link.get("href").find("http") != -1:
                news_link[link.get_text()] = link.get('href')
                # print(link.get_text())
    return news_link

def page_url(url, page_num):#生成带页面的URL
    if page_num == 1:
        return url
    index = url.rfind("/")
    return url[0 : index+1] + "index_" + str(page_num) + '.html'

def get_news_body(url):#抓取新闻主体内容
    first = True
    content_text = []
    page_num = 1
    article_div = ""

    #使用循环处理有分页的新闻
    # while first == True or article_div.find("下一页</a>") != -1:
    soup = get_html_soup(url, 'gbk')
    if soup == None:
        return None
    article_div = str(soup.find("div", attrs = {"id": "article"}))
    soup = BeautifulSoup(str(article_div), "lxml")
    para_arr = soup.find_all("p")
    lenth = len(para_arr)
    for i in range(0,lenth - 2):
        if len(para_arr[i].get_text().strip()) > 0:
            content_text.append("    " + para_arr[i].get_text().strip())
    for x in content_text:
        if x == "    None":
            return None
    return content_text

def clean_chinese_character(text):
    '''处理特殊的中文符号,将其全部替换为'-' 否则在保存时Windows无法将有的中文符号作为路径'''
    chars = chars = ["/", "\"", "'", "·", "。","？", "！", "，", "、", "；", "：", "‘", "’", "“", "”", "（", "）", "…", "–", "．", "《", "》","|","?",","];
    new_text = ""
    for i in range(len(text)):
        if text[i] not in chars:
            new_text += text[i]
        else:
            new_text += "_"
    return new_text;

def create_docx(news_type, title, content):
    #'''这里使用python-docx库将新闻的内容生成word文件'''
    document = Document()
    paragraph = document.add_paragraph(title)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.bold = True

    for x in content:
        paragraph = document.add_paragraph(x)

    style = paragraph.style
    font = style.font
    font.size = Pt(15)
    font.name = "consolas"

    name = clean_chinese_character(title) + ".docx"
    document.save(news_type + "/" + name)
    
def create_txt(news_type, title, content):
    #'''这里使用python-docx库将新闻的内容生成word文件'''
    # document = Document()
    # paragraph = document.add_paragraph(title)
    # paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # paragraph.bold = True

    # for x in content:
        # paragraph = document.add_paragraph(x)

    # style = paragraph.style
    # font = style.font
    # font.size = Pt(15)
    # font.name = "consolas"
    name=news_type+'/'
    name+= clean_chinese_character(title) + ".txt"
    f=codecs.open(name,'w','utf-8')
    for x in content:
        f.write(x)
        f.write('\n')
    
########################################################################
origin_url="http://www.yangtse.com"
origin_pattern = {"class": "left"}
news="F:/news"

print("deleting old dir")
if os.path.exists(news):
    shutil.rmtree(news,True)

print("creating dir: ", news)
os.mkdir(news)
class_list=getClasslink(origin_url, origin_pattern)
print("\ngetting news")

for x in class_list:
    class_url=class_list[x];
    arr=class_url.split('/');
    type_len=len(arr)
    news_type=arr[type_len-2]
    title=news+'/'+news_type
    if not os.path.exists(title):
        os.mkdir(title)
    news_list=get_title_link(class_url,{"class": "nodelist"},10);
    for y in news_list:
        paras = get_news_body(news_list[y])
        if paras != None and len(paras) > 0:
            print("writing:")
            create_txt(title, y, paras)

print("All done, have a nice day")


    
    
    


